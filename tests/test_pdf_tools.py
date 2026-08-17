from pathlib import Path
import io
import unittest
from unittest.mock import Mock, patch

from PIL import Image
from pypdf import PdfReader, PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from app import copy_pages, parse_pages
from pdf_core import (
    BookmarkItem,
    MarkupAnnotation,
    PageItem,
    TextBlock,
    add_bates_numbering,
    add_page_numbers,
    apply_markup_annotations,
    apply_outline,
    apply_text_markups_for_query,
    build_markup_annotation,
    compare_pdf_text,
    compare_pdf_visual,
    crop_pdf_pages,
    encrypt_pdf_with_permissions,
    extract_outline,
    extract_page_text_blocks,
    extract_pdf_page_indexes,
    apply_edits_and_extract_pages,
    resolve_system_font_file,
    fill_form_fields,
    flatten_form_fields,
    images_to_pdf,
    insert_pdf_pages,
    office_app_for_path,
    office_files_to_pdf,
    convert_office_file_to_pdf,
    _libreoffice_pdf_filter,
    suggested_pdf_name_for_source,
    suggested_pdf_path_for_source,
    _hide_com_window,
    list_form_fields,
    parse_page_groups,
    replace_pdf_pages,
    secure_redact_query,
    split_pdf_advanced,
    split_pdf_by_bookmarks,
    add_text_overlay_annotation,
    add_text_overlay_annotations,
    resolve_annotation_fill,
    iter_markup_pdf_annotations,
    add_text_stamp,
    add_watermark,
    build_annotation_da,
    overlay_needs_embedded_font,
    text_contains_cjk,
    callout_layout,
    callout_box_from_pointer,
    comment_box_from_pointer,
    comment_polyline,
    clean_metadata,
    merge_pdf_files,
    merge_text_blocks,
    ocr_pdf_to_searchable_pdf,
    ocr_pdf_to_text,
    detect_ocr_language,
    last_ocr_language,
    ocr_language_short_label,
    _docx_body_font_name,
    _is_noisy_ocr_token,
    _ocr_words_to_plain_text,
    _ocr_words_to_rows,
    _detect_page_gutter_x,
    _add_docx_ocr_page,
    _add_docx_layout_table,
    _configure_docx_section_from_pdf,
    _page_rows_are_two_column,
    _rows_are_side_by_side_layout,
    _column_tess_lang,
    _tesseract_column_config,
    configure_tesseract,
    ensure_ocr_available,
    normalize_cjk_text,
    count_cjk_chars,
    pdf_to_docx,
    pdf_to_images,
    pdf_to_xlsx,
    _is_useful_table,
    _pymupdf_page_tables,
    _table_looks_fragmented,
    _words_to_table_rows,
    redact_matching_text_blocks_overlay,
    redact_text_block_secure,
    redact_text_block_overlay,
    replace_text_block_content_stream,
    replace_text_block_overlay,
    replace_text_block_seamless,
    remove_blank_pages,
    split_pdf_to_zip,
    text_matches_query,
    EraseMark,
    apply_erase_marks,
    apply_erase_then_text_overlays,
    write_page_items_merged,
    write_page_items_separately,
)


class PdfToolsTests(unittest.TestCase):
    def test_parse_pages_keeps_order_and_supports_ranges(self):
        self.assertEqual(parse_pages("1,3,5-7", 10), [0, 2, 4, 5, 6])

    def test_parse_pages_removes_duplicates(self):
        self.assertEqual(parse_pages("1,1,2-3,3", 4), [0, 1, 2])

    def test_copy_pages(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=100, height=100)
        writer.add_blank_page(width=100, height=100)
        with source.open("wb") as stream:
            writer.write(stream)

        reader = PdfReader(str(source))
        copied = copy_pages(reader, [1])
        target = Path(self.temp_dir.name) / "target.pdf"
        with target.open("wb") as stream:
            copied.write(stream)

        self.assertEqual(len(PdfReader(str(target)).pages), 1)

    def test_extract_pdf_page_indexes_keeps_selected_order(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "target.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=100, height=100)
        writer.add_blank_page(width=120, height=120)
        writer.add_blank_page(width=140, height=140)
        with source.open("wb") as stream:
            writer.write(stream)

        count = extract_pdf_page_indexes(source, target, [2, 0, 2, -1, 9])
        self.assertEqual(count, 2)
        pages = PdfReader(str(target)).pages
        self.assertEqual(len(pages), 2)
        self.assertAlmostEqual(float(pages[0].mediabox.width), 140)
        self.assertAlmostEqual(float(pages[1].mediabox.width), 100)

    def test_apply_edits_and_extract_pages_keeps_only_annotated_page(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "target.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=400)
        writer.add_blank_page(width=300, height=400)
        writer.add_blank_page(width=300, height=400)
        with source.open("wb") as stream:
            writer.write(stream)

        count = apply_edits_and_extract_pages(
            source,
            target,
            [0],
            overlays=[
                {
                    "page_index": 0,
                    "text": "Note here",
                    "font_size": 12,
                    "cover": False,
                    "font_key": "helvetica",
                    "bold": False,
                    "color_rgb": (0.0, 0.0, 0.0),
                    "pdf_x": 72,
                    "pdf_y": 300,
                    "rect_width": 160,
                    "rect_height": 32,
                    "shape": "box",
                }
            ],
        )
        self.assertEqual(count, 1)
        reader = PdfReader(str(target))
        self.assertEqual(len(reader.pages), 1)
        annots = reader.pages[0].get("/Annots")
        self.assertTrue(annots)

    def test_build_annotation_da_supports_bold_and_color(self):
        self.assertEqual(
            build_annotation_da("helvetica", 14, True, (0.8, 0.0, 0.0)),
            "/Helv-Bold 14 Tf 0.8 0.0 0.0 rg",
        )

    def test_add_text_overlay_annotation(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "target.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=400)
        with source.open("wb") as stream:
            writer.write(stream)

        add_text_overlay_annotation(
            source=source,
            target=target,
            page_index=0,
            x=72,
            y=300,
            text="Reviewed FY2026",
            font_size=12,
            cover_original=True,
            cover_width=180,
            cover_height=30,
        )

        annots = PdfReader(str(target)).pages[0].get("/Annots")
        self.assertEqual(len(annots), 2)

    def test_add_text_overlay_annotations_applies_multiple_pages(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "target.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=400)
        writer.add_blank_page(width=300, height=400)
        with source.open("wb") as stream:
            writer.write(stream)

        overlay = {
            "pdf_x": 72,
            "pdf_y": 300,
            "font_size": 12,
            "cover": True,
            "rect_width": 180,
            "rect_height": 30,
            "font_key": "helvetica",
            "bold": False,
            "color_rgb": (0.0, 0.0, 0.0),
            "shape": "box",
        }
        add_text_overlay_annotations(
            source,
            target,
            [
                {**overlay, "page_index": 0, "text": "First"},
                {**overlay, "page_index": 1, "text": "Second", "pdf_y": 280},
            ],
        )

        reader = PdfReader(str(target))
        self.assertEqual(len(reader.pages[0].get("/Annots")), 2)
        self.assertEqual(len(reader.pages[1].get("/Annots")), 2)

    def test_add_text_overlay_speech_shape_writes_polygon(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "target.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=400)
        with source.open("wb") as stream:
            writer.write(stream)

        add_text_overlay_annotation(
            source=source,
            target=target,
            page_index=0,
            x=80,
            y=220,
            text="看這裡",
            font_size=12,
            cover_original=True,
            cover_width=160,
            cover_height=36,
            shape="speech",
            pointer=(40, 180),
        )

        annots = PdfReader(str(target)).pages[0].get("/Annots")
        self.assertGreaterEqual(len(annots), 2)
        subtypes = [str(item.get_object().get("/Subtype")) for item in annots]
        self.assertIn("/Polygon", subtypes)
        self.assertIn("/FreeText", subtypes)

    def test_add_text_overlay_rect_shape_writes_frame(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "target.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=400)
        with source.open("wb") as stream:
            writer.write(stream)

        add_text_overlay_annotation(
            source=source,
            target=target,
            page_index=0,
            x=80,
            y=220,
            text="框住這段",
            font_size=12,
            cover_original=True,
            cover_width=160,
            cover_height=36,
            shape="rect",
        )

        annots = PdfReader(str(target)).pages[0].get("/Annots")
        subtypes = [str(item.get_object().get("/Subtype")) for item in annots]
        self.assertIn("/Square", subtypes)

    def test_text_contains_cjk(self):
        self.assertTrue(text_contains_cjk("這個要修改"))
        self.assertTrue(overlay_needs_embedded_font("Hello", "cjk"))
        self.assertTrue(overlay_needs_embedded_font("這個要修改", "courier"))
        self.assertFalse(text_contains_cjk("Reviewed FY2026"))
        self.assertFalse(overlay_needs_embedded_font("Reviewed FY2026", "courier"))

    def test_add_text_overlay_embeds_chinese_even_with_courier(self):
        import fitz

        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "target.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=400)
        with source.open("wb") as stream:
            writer.write(stream)

        add_text_overlay_annotation(
            source=source,
            target=target,
            page_index=0,
            x=72,
            y=300,
            text="這個要修改",
            font_size=18,
            cover_original=False,
            cover_width=220,
            cover_height=32,
            font_key="courier",
        )

        document = fitz.open(str(target))
        try:
            self.assertIn("這個要修改", document[0].get_text())
        finally:
            document.close()
        self.assertLess(target.stat().st_size, 200_000)

    def test_callout_layout_uses_explicit_box_size(self):
        layout = callout_layout(
            MarkupAnnotation(
                "callout",
                40,
                50,
                100,
                200,
                contents="註解",
                box_width=180,
                box_height=40,
            )
        )
        self.assertAlmostEqual(layout.left, 100)
        self.assertAlmostEqual(layout.bottom, 200)
        self.assertAlmostEqual(layout.right, 280)
        self.assertAlmostEqual(layout.top, 240)
        self.assertEqual(layout.pointer, (40, 50))

    def test_callout_box_from_pointer_centers_box_above_arrow(self):
        left, bottom = callout_box_from_pointer(100, 50, 180, 40, gap=18)
        self.assertAlmostEqual(left, 10)
        self.assertAlmostEqual(bottom, 68)

    def test_comment_box_from_pointer_offsets_box_to_the_side(self):
        left, bottom = comment_box_from_pointer(100, 50, 180, 40)
        self.assertAlmostEqual(left, 90)
        self.assertAlmostEqual(bottom, 78)
        self.assertGreater(left + 90, 100)

    def test_comment_polyline_has_stretchable_horizontal_arm(self):
        path = comment_polyline(90, 78, 270, 118, 100, 50)
        self.assertEqual(len(path), 4)
        self.assertAlmostEqual(path[0][0], path[1][0])
        self.assertAlmostEqual(path[1][1], path[2][1])
        self.assertAlmostEqual(path[2][0], path[3][0])
        self.assertEqual(path[3], (100, 50))
        self.assertGreater(abs(path[2][0] - path[1][0]), 40)

    def test_add_text_overlay_comment_writes_polyline(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "target.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=400)
        with source.open("wb") as stream:
            writer.write(stream)

        add_text_overlay_annotation(
            source=source,
            target=target,
            page_index=0,
            x=90,
            y=220,
            text="Note here",
            font_size=12,
            cover_original=False,
            cover_width=160,
            cover_height=36,
            shape="comment",
            pointer=(40, 180),
        )

        annots = PdfReader(str(target)).pages[0].get("/Annots")
        subtypes = [str(item.get_object().get("/Subtype")) for item in annots]
        self.assertIn("/Square", subtypes)
        self.assertIn("/PolyLine", subtypes)
        self.assertIn("/FreeText", subtypes)
        polyline = next(item.get_object() for item in annots if str(item.get_object().get("/Subtype")) == "/PolyLine")
        vertices = [float(value) for value in polyline.get("/Vertices")]
        self.assertEqual(len(vertices), 8)

    def test_resolve_annotation_fill_supports_none_and_custom(self):
        self.assertIsNone(resolve_annotation_fill((0, 0, 0), fill_none=True))
        self.assertEqual(resolve_annotation_fill((0, 0, 0), fill_rgb=(1.0, 0.96, 0.72)), (1.0, 0.96, 0.72))
        auto = resolve_annotation_fill((0.0, 0.0, 0.0))
        self.assertGreater(auto[0], 0.7)

    def test_comment_square_uses_custom_fill_color(self):
        markup = MarkupAnnotation(
            "comment",
            40,
            50,
            100,
            200,
            color_rgb=(0.0, 0.0, 0.0),
            contents="Note",
            box_width=160,
            box_height=36,
            fill_rgb=(1.0, 0.96, 0.72),
        )
        square = next(
            annotation
            for annotation in iter_markup_pdf_annotations(markup)
            if str(annotation.get("/Subtype")) == "/Square"
        )
        fill = [float(value) for value in square.get("/IC")]
        self.assertAlmostEqual(fill[0], 1.0)
        self.assertAlmostEqual(fill[1], 0.96)
        self.assertAlmostEqual(fill[2], 0.72)

    def test_comment_square_omits_fill_when_none(self):
        markup = MarkupAnnotation(
            "comment",
            40,
            50,
            100,
            200,
            color_rgb=(0.0, 0.0, 0.0),
            contents="Note",
            box_width=160,
            box_height=36,
            fill_none=True,
        )
        square = next(
            annotation
            for annotation in iter_markup_pdf_annotations(markup)
            if str(annotation.get("/Subtype")) == "/Square"
        )
        self.assertIsNone(square.get("/IC"))

    def test_office_app_for_path_detects_word_excel_ppt(self):
        self.assertEqual(office_app_for_path(Path("memo.docx")), "word")
        self.assertEqual(office_app_for_path(Path("book.xlsx")), "excel")
        self.assertEqual(office_app_for_path(Path("deck.pptx")), "powerpoint")
        with self.assertRaises(ValueError):
            office_app_for_path(Path("notes.txt"))

    def test_libreoffice_pdf_filter_keeps_image_and_font_settings(self):
        word_filter = _libreoffice_pdf_filter(Path("memo.docx"))
        excel_filter = _libreoffice_pdf_filter(Path("book.xlsx"))
        ppt_filter = _libreoffice_pdf_filter(Path("deck.pptx"))
        self.assertIn("writer_pdf_Export", word_filter)
        self.assertIn("calc_pdf_Export", excel_filter)
        self.assertIn("impress_pdf_Export", ppt_filter)
        self.assertIn("UseLosslessCompression", word_filter)
        self.assertIn("ReduceImageResolution", word_filter)
        self.assertIn("EmbedStandardFonts", word_filter)
        self.assertIn('"value":"false"', word_filter)

    def test_convert_office_file_requires_backend(self):
        source = Path(self.temp_dir.name) / "memo.docx"
        target = Path(self.temp_dir.name) / "memo.pdf"
        source.write_bytes(b"fake-docx")
        with patch("pdf_core._convert_office_with_com", return_value=(False, "COM 失敗")):
            with patch("pdf_core.find_libreoffice_executable", return_value=None):
                with self.assertRaises(ValueError) as ctx:
                    convert_office_file_to_pdf(source, target)
        message = str(ctx.exception)
        self.assertIn("Office", message)
        self.assertIn("COM 失敗", message)

    def test_convert_office_file_uses_com_success(self):
        source = Path(self.temp_dir.name) / "memo.docx"
        target = Path(self.temp_dir.name) / "memo.pdf"
        source.write_bytes(b"fake-docx")

        def fake_com(src: Path, dest: Path, progress=None, pump=None):
            seen.append(src)
            writer = PdfWriter()
            writer.add_blank_page(width=300, height=400)
            with dest.open("wb") as stream:
                writer.write(stream)
            return True, ""

        seen: list[Path] = []
        with patch("pdf_core._convert_office_with_com", side_effect=fake_com):
            with patch("pdf_core.find_libreoffice_executable") as find_lo:
                convert_office_file_to_pdf(source, target)
        find_lo.assert_not_called()
        self.assertTrue(target.exists())
        self.assertEqual(seen[0].name, source.name)
        self.assertNotEqual(seen[0], source)

    def test_suggested_pdf_name_keeps_chinese_stem(self):
        source = Path(r"T:\Reporting\中海宏洋2026年中期业绩简报 v.1.pptx")
        self.assertEqual(
            suggested_pdf_name_for_source(source),
            "中海宏洋2026年中期业绩简报 v.1.pdf",
        )
        self.assertEqual(
            suggested_pdf_path_for_source(source),
            Path(r"T:\Reporting\中海宏洋2026年中期业绩简报 v.1.pdf"),
        )

    def test_hide_com_window_prefers_invisible_app(self):
        class FakeApp:
            Visible = True
            DisplayAlerts = 2
            HWND = 0

        app = FakeApp()
        _hide_com_window(app)
        self.assertFalse(app.Visible)

    def test_images_to_pdf_lossless_keeps_page_size(self):
        image_path = Path(self.temp_dir.name) / "slide.png"
        Image.new("RGB", (3000, 1688), (240, 240, 240)).save(image_path)
        target = Path(self.temp_dir.name) / "slide.pdf"
        images_to_pdf([image_path], target, resolution=300, lossless=True)
        page = PdfReader(str(target)).pages[0]
        self.assertAlmostEqual(float(page.mediabox.width), 720.0, delta=2)
        self.assertAlmostEqual(float(page.mediabox.height), 405.12, delta=2)

    def test_office_files_to_pdf_merges_converted_files(self):
        source_a = Path(self.temp_dir.name) / "a.docx"
        source_b = Path(self.temp_dir.name) / "b.xlsx"
        source_a.write_bytes(b"fake-a")
        source_b.write_bytes(b"fake-b")
        target = Path(self.temp_dir.name) / "combined.pdf"

        def fake_convert(source: Path, dest: Path, progress=None, pump=None) -> None:
            writer = PdfWriter()
            writer.add_blank_page(width=300, height=400)
            dest.parent.mkdir(parents=True, exist_ok=True)
            with dest.open("wb") as stream:
                writer.write(stream)

        with patch("pdf_core.convert_office_file_to_pdf", side_effect=fake_convert):
            count = office_files_to_pdf([source_a, source_b], target)

        self.assertEqual(count, 2)
        self.assertEqual(len(PdfReader(str(target)).pages), 2)

    def test_replace_text_block_overlay(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "target.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=400)
        with source.open("wb") as stream:
            writer.write(stream)

        replace_text_block_overlay(
            source,
            target,
            0,
            TextBlock("Old", 72, 300, 80, 24, 12, "/Helvetica"),
            "New",
        )

        annots = PdfReader(str(target)).pages[0].get("/Annots")
        self.assertEqual(len(annots), 2)

    def test_redact_text_block_overlay_adds_black_box(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "target.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=400)
        with source.open("wb") as stream:
            writer.write(stream)

        redact_text_block_overlay(
            source,
            target,
            0,
            TextBlock("Secret", 72, 300, 80, 24, 12, "/Helvetica"),
        )

        annot = PdfReader(str(target)).pages[0].get("/Annots")[0].get_object()
        self.assertEqual(annot.get("/Subtype"), "/Square")
        self.assertEqual(list(annot.get("/IC")), [0, 0, 0])

    def test_redact_matching_text_blocks_overlay_adds_all_matches(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "target.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=400)
        writer.add_blank_page(width=300, height=400)
        with source.open("wb") as stream:
            writer.write(stream)
        page_blocks = {
            0: [TextBlock("Invoice Number", 72, 300, 80, 24, 12, "/Helvetica")],
            1: [TextBlock("Invoice Total", 72, 260, 80, 24, 12, "/Helvetica")],
        }

        with patch("pdf_core.extract_page_text_blocks", side_effect=lambda _source, page, _password="": page_blocks[page]):
            count = redact_matching_text_blocks_overlay(source, target, "invoice")

        reader = PdfReader(str(target))
        self.assertEqual(count, 2)
        self.assertEqual(len(reader.pages[0].get("/Annots")), 1)
        self.assertEqual(len(reader.pages[1].get("/Annots")), 1)

    def test_text_matches_query_supports_case_and_whole_word(self):
        self.assertTrue(text_matches_query("Invoice Number", "invoice"))
        self.assertFalse(text_matches_query("Invoice Number", "invoice", case_sensitive=True))
        self.assertTrue(text_matches_query("Invoice Number", "Invoice", case_sensitive=True))
        self.assertTrue(text_matches_query("Invoice Number", "Invoice", whole_word=True))
        self.assertFalse(text_matches_query("InvoiceNumber", "Invoice", whole_word=True))

    def test_redact_text_block_secure_removes_simple_text_and_adds_black_box(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "target.pdf"
        writer = PdfWriter()
        page = writer.add_blank_page(width=300, height=400)
        stream = DecodedStreamObject()
        stream.set_data(b"BT /F1 12 Tf 72 300 Td (Secret) Tj ET")
        page[NameObject("/Contents")] = stream
        page[NameObject("/Resources")] = DictionaryObject()
        with source.open("wb") as output:
            writer.write(output)

        redact_text_block_secure(
            source,
            target,
            0,
            TextBlock("Secret", 72, 300, 80, 24, 12, "/Helvetica"),
        )

        page = PdfReader(str(target)).pages[0]
        self.assertNotIn(b"Secret", page.get_contents().get_data())
        self.assertEqual(len(page.get("/Annots")), 1)

    def test_replace_text_block_content_stream_rewrites_simple_text(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "target.pdf"
        writer = PdfWriter()
        page = writer.add_blank_page(width=300, height=400)
        stream = DecodedStreamObject()
        stream.set_data(b"BT /F1 12 Tf 72 300 Td (Old123) Tj ET")
        page[NameObject("/Contents")] = stream
        page[NameObject("/Resources")] = DictionaryObject()
        with source.open("wb") as output:
            writer.write(output)

        replace_text_block_content_stream(
            source,
            target,
            0,
            TextBlock("Old123", 72, 300, 80, 24, 12, "/Helvetica"),
            "New456",
        )

        content = PdfReader(str(target)).pages[0].get_contents().get_data()
        self.assertIn(b"New456", content)
        self.assertNotIn(b"Old123", content)

    def test_replace_text_block_content_stream_requires_same_length(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "target.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=400)
        with source.open("wb") as output:
            writer.write(output)

        with self.assertRaises(ValueError):
            replace_text_block_content_stream(
                source,
                target,
                0,
                TextBlock("Old123", 72, 300, 80, 24, 12, "/Helvetica"),
                "Longer text",
            )

    def test_merge_text_blocks_combines_same_line_fragments(self):
        blocks = [
            TextBlock("Hello", 10, 100, 30, 12, 10),
            TextBlock("World", 45, 101, 30, 12, 10),
            TextBlock("Next", 10, 70, 24, 12, 10),
        ]

        merged = merge_text_blocks(blocks)

        self.assertEqual([block.text for block in merged], ["Hello World", "Next"])

    def test_add_page_numbers(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "target.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=400)
        writer.add_blank_page(width=300, height=400)
        with source.open("wb") as stream:
            writer.write(stream)

        add_page_numbers(source, target, "Page {page} of {total}")

        reader = PdfReader(str(target))
        self.assertEqual(len(reader.pages[0].get("/Annots")), 1)
        self.assertEqual(len(reader.pages[1].get("/Annots")), 1)

    def test_add_watermark(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "target.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=400)
        with source.open("wb") as stream:
            writer.write(stream)

        add_watermark(source, target, "CONFIDENTIAL")

        import fitz

        document = fitz.open(str(target))
        try:
            self.assertIn("CONFIDENTIAL", document[0].get_text())
        finally:
            document.close()

    def test_add_watermark_respects_position_and_custom_text(self):
        source = Path(self.temp_dir.name) / "wm-source.pdf"
        top_right = Path(self.temp_dir.name) / "wm-top-right.pdf"
        bottom_left = Path(self.temp_dir.name) / "wm-bottom-left.pdf"
        custom = Path(self.temp_dir.name) / "wm-custom.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=400)
        writer.add_blank_page(width=300, height=400)
        with source.open("wb") as stream:
            writer.write(stream)

        add_watermark(source, top_right, "DRAFT", position="top-right", rotation=0, font_size=18, opacity=1.0)
        add_watermark(source, bottom_left, "DRAFT", position="bottom-left", rotation=0, font_size=18, opacity=1.0)
        add_watermark(
            source,
            custom,
            "機密",
            position="custom",
            rotation=0,
            font_size=16,
            opacity=1.0,
            x_percent=15,
            y_percent=80,
            pages_spec="1",
        )

        import fitz

        def first_word_box(path: Path):
            document = fitz.open(str(path))
            try:
                words = document[0].get_text("words")
                self.assertTrue(words)
                return words[0][:4], document[0].rect
            finally:
                document.close()

        (x0, y0, x1, y1), rect = first_word_box(top_right)
        self.assertGreater(x0, rect.width / 2)
        self.assertLess(y1, rect.height / 2)

        (x0, y0, x1, y1), rect = first_word_box(bottom_left)
        self.assertLess(x1, rect.width / 2)
        self.assertGreater(y0, rect.height / 2)

        document = fitz.open(str(custom))
        try:
            self.assertIn("機密", document[0].get_text())
            self.assertEqual(document[1].get_text().strip(), "")
            words = document[0].get_text("words")
            x0, y0, x1, y1, *_rest = words[0]
            self.assertLess(x1, document[0].rect.width / 2)
            self.assertGreater(y0, document[0].rect.height / 2)
        finally:
            document.close()

        with self.assertRaises(ValueError):
            add_watermark(source, Path(self.temp_dir.name) / "empty.pdf", "   ")

    def test_clean_metadata(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "target.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=400)
        writer.add_metadata({"/Author": "Sensitive User"})
        with source.open("wb") as stream:
            writer.write(stream)

        clean_metadata(source, target)

        metadata = PdfReader(str(target)).metadata
        self.assertNotEqual(metadata.get("/Author"), "Sensitive User")

    def test_remove_blank_pages_refuses_all_blank_output(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "target.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=400)
        with source.open("wb") as stream:
            writer.write(stream)

        with self.assertRaises(ValueError):
            remove_blank_pages(source, target)

    def test_write_page_items_merged(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "target.pdf"
        writer = PdfWriter()
        for _ in range(3):
            writer.add_blank_page(width=300, height=400)
        with source.open("wb") as stream:
            writer.write(stream)

        page_items = [PageItem(source, index, f"Page {index + 1}") for index in range(3)]
        write_page_items_merged(page_items, [0, 2], target)

        self.assertEqual(len(PdfReader(str(target)).pages), 2)

    def test_write_page_items_merged_applies_rotation(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "target.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=400)
        with source.open("wb") as stream:
            writer.write(stream)

        page_items = [PageItem(source, 0, "Page 1", rotation=90)]
        write_page_items_merged(page_items, [0], target)

        self.assertEqual(PdfReader(str(target)).pages[0].get("/Rotate"), 90)

    def test_merge_pdf_files(self):
        first = Path(self.temp_dir.name) / "first.pdf"
        second = Path(self.temp_dir.name) / "second.pdf"
        target = Path(self.temp_dir.name) / "merged.pdf"
        for path in (first, second):
            writer = PdfWriter()
            writer.add_blank_page(width=300, height=400)
            with path.open("wb") as stream:
                writer.write(stream)

        merge_pdf_files([first, second], target)

        self.assertEqual(len(PdfReader(str(target)).pages), 2)

    def test_pdf_to_images_writes_png_files(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        folder = Path(self.temp_dir.name) / "images"
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=400)
        writer.add_blank_page(width=300, height=400)
        with source.open("wb") as stream:
            writer.write(stream)

        count = pdf_to_images(source, folder, image_format="png", dpi=72)

        self.assertEqual(count, 2)
        self.assertEqual(len(list(folder.glob("*.png"))), 2)

    def test_pdf_to_images_supports_page_ranges(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        folder = Path(self.temp_dir.name) / "images-range"
        writer = PdfWriter()
        for _ in range(3):
            writer.add_blank_page(width=300, height=400)
        with source.open("wb") as stream:
            writer.write(stream)

        count = pdf_to_images(source, folder, image_format="jpg", dpi=72, pages_spec="1,3")

        self.assertEqual(count, 2)
        self.assertEqual(len(list(folder.glob("*.jpg"))), 2)

    def test_ocr_pdf_to_text_writes_text(self):
        target = Path(self.temp_dir.name) / "ocr.txt"
        with patch("pdf_core.ensure_ocr_available"), patch(
            "pdf_core.iter_pdf_page_images",
            return_value=[(0, Image.new("RGB", (10, 10), "white"))],
        ), patch("pdf_core.pytesseract", Mock(image_to_string=Mock(return_value="Hello OCR"))):
            count = ocr_pdf_to_text(Path("source.pdf"), target, language="eng")

        self.assertEqual(count, 1)
        self.assertIn("Hello OCR", target.read_text(encoding="utf-8"))

    def test_normalize_cjk_text_fixes_radicals_and_spaces(self):
        self.assertEqual(normalize_cjk_text("中 英 ⽂ ⾴ ⾦"), "中英文頁金")
        self.assertGreaterEqual(count_cjk_chars("金額核對"), 4)

    def test_ocr_pdf_to_text_uses_chinese_primary_language_and_normalizes(self):
        target = Path(self.temp_dir.name) / "ocr-cjk.txt"
        ocr = Mock(image_to_string=Mock(return_value="中 英 ⽂ 年 報"))
        with patch("pdf_core.ensure_ocr_available"), patch(
            "pdf_core.iter_pdf_page_images",
            return_value=[(0, Image.new("RGB", (10, 10), "white"))],
        ), patch("pdf_core.pytesseract", ocr):
            count = ocr_pdf_to_text(Path("source.pdf"), target, language="eng+chi_tra", dpi=200)

        self.assertEqual(count, 1)
        self.assertIn("中英文年報", target.read_text(encoding="utf-8"))
        kwargs = ocr.image_to_string.call_args.kwargs
        self.assertEqual(kwargs["lang"], "chi_tra+eng")
        self.assertIn("--psm 3", kwargs["config"])

    def test_detect_ocr_language_simplified_traditional_and_english(self):
        self.assertEqual(
            detect_ocr_language("深圳投资审计报告会计师事务所财务报表"),
            "eng+chi_sim",
        )
        self.assertEqual(
            detect_ocr_language("這份報告會與財務報表核對並請審閱"),
            "eng+chi_tra",
        )
        self.assertEqual(
            detect_ocr_language("Quarterly report for ABC Limited"),
            "eng",
        )
        self.assertEqual(ocr_language_short_label("eng+chi_sim"), "英文+簡中")

    def test_ocr_pdf_to_text_auto_detects_simplified_chinese(self):
        target = Path(self.temp_dir.name) / "ocr-auto.txt"
        ocr = Mock(image_to_string=Mock(return_value="深圳投资审计报告会计师事务所"))
        with patch("pdf_core.ensure_ocr_available"), patch(
            "pdf_core.iter_pdf_page_images",
            return_value=[(0, Image.new("RGB", (10, 10), "white"))],
        ), patch("pdf_core.pytesseract", ocr):
            count = ocr_pdf_to_text(Path("source.pdf"), target, language="auto", dpi=200)

        self.assertEqual(count, 1)
        self.assertEqual(last_ocr_language(), "eng+chi_sim")
        self.assertEqual(ocr.image_to_string.call_args.kwargs["lang"], "chi_sim+eng")
        self.assertIn("深圳投资审计报告", target.read_text(encoding="utf-8"))

    def test_ocr_pdf_to_text_reports_page_progress(self):
        target = Path(self.temp_dir.name) / "ocr-progress.txt"
        seen: list[str] = []
        with patch("pdf_core.ensure_ocr_available"), patch(
            "pdf_core.iter_pdf_page_images",
            return_value=[
                (0, Image.new("RGB", (10, 10), "white")),
                (1, Image.new("RGB", (10, 10), "white")),
            ],
        ), patch("pdf_core.pytesseract", Mock(image_to_string=Mock(return_value="page"))):
            count = ocr_pdf_to_text(
                Path("source.pdf"),
                target,
                language="eng",
                progress=lambda current, total, text: seen.append(text),
            )
        self.assertEqual(count, 2)
        self.assertTrue(any("第 1 頁" in text for text in seen))
        self.assertTrue(any("第 2 頁" in text for text in seen))

    def test_ocr_pdf_to_searchable_pdf_merges_ocr_pages(self):
        target = Path(self.temp_dir.name) / "searchable.pdf"
        pdf_stream = io.BytesIO()
        writer = PdfWriter()
        writer.add_blank_page(width=100, height=100)
        writer.write(pdf_stream)
        with patch("pdf_core.ensure_ocr_available"), patch(
            "pdf_core.iter_pdf_page_images",
            return_value=[(0, Image.new("RGB", (10, 10), "white"))],
        ), patch(
            "pdf_core.pytesseract",
            Mock(image_to_pdf_or_hocr=Mock(return_value=pdf_stream.getvalue())),
        ):
            count = ocr_pdf_to_searchable_pdf(Path("source.pdf"), target, language="eng")

        self.assertEqual(count, 1)
        self.assertEqual(len(PdfReader(str(target)).pages), 1)

    def test_configure_tesseract_uses_discovered_exe(self):
        fake_dir = Path(self.temp_dir.name) / "Tesseract-OCR"
        fake_dir.mkdir()
        fake_exe = fake_dir / "tesseract.exe"
        fake_exe.write_bytes(b"stub")
        mock = Mock()
        mock.get_tesseract_version.side_effect = [Exception("not on PATH"), "5.4.0"]
        mock.pytesseract = Mock()
        with patch("pdf_core.OCR_AVAILABLE", True), patch("pdf_core.pytesseract", mock):
            with patch("pdf_core._tesseract_exe_candidates", return_value=[fake_exe]):
                found = configure_tesseract()
        self.assertEqual(found, fake_exe)
        self.assertEqual(mock.pytesseract.tesseract_cmd, str(fake_exe))

    def test_ensure_ocr_available_explains_scan_pdf_when_missing(self):
        with patch("pdf_core.OCR_AVAILABLE", False), patch("pdf_core.pytesseract", None):
            with self.assertRaises(ValueError) as raised:
                ensure_ocr_available()
        self.assertIn("掃描件", str(raised.exception))
        self.assertIn("Tesseract", str(raised.exception))

    def test_split_pdf_to_zip(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "split.zip"
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=400)
        writer.add_blank_page(width=300, height=400)
        with source.open("wb") as stream:
            writer.write(stream)

        page_count = split_pdf_to_zip(source, target)

        self.assertEqual(page_count, 2)
        self.assertTrue(target.exists())

    def test_write_page_items_separately(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        folder = Path(self.temp_dir.name) / "out"
        writer = PdfWriter()
        for _ in range(3):
            writer.add_blank_page(width=300, height=400)
        with source.open("wb") as stream:
            writer.write(stream)

        page_items = [PageItem(source, index, f"Page {index + 1}") for index in range(3)]
        count = write_page_items_separately(page_items, [0, 2], folder)

        self.assertEqual(count, 2)
        self.assertEqual(len(list(folder.glob("*.pdf"))), 2)

    def test_apply_and_extract_outline_round_trip(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "target.pdf"
        writer = PdfWriter()
        for _ in range(5):
            writer.add_blank_page(width=300, height=400)
        with source.open("wb") as stream:
            writer.write(stream)

        items = [
            BookmarkItem("封面", 0, 0),
            BookmarkItem("第一章", 1, 0),
            BookmarkItem("第一節", 2, 1),
            BookmarkItem("第二章", 3, 0),
        ]
        apply_outline(source, target, items)

        extracted = extract_outline(target)
        self.assertEqual(
            [(item.title, item.page_index, item.level) for item in extracted],
            [("封面", 0, 0), ("第一章", 1, 0), ("第一節", 2, 1), ("第二章", 3, 0)],
        )

    def test_apply_outline_clamps_out_of_range_pages(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "target.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=400)
        writer.add_blank_page(width=300, height=400)
        with source.open("wb") as stream:
            writer.write(stream)

        apply_outline(source, target, [BookmarkItem("超出範圍", 99, 0)])

        extracted = extract_outline(target)
        self.assertEqual(len(extracted), 1)
        self.assertEqual(extracted[0].page_index, 1)

    def test_extract_outline_returns_empty_without_bookmarks(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=400)
        with source.open("wb") as stream:
            writer.write(stream)

        self.assertEqual(extract_outline(source), [])

    def test_build_markup_annotation_highlight_has_quadpoints(self):
        annotation = build_markup_annotation(
            MarkupAnnotation("highlight", 10, 20, 110, 40, (1.0, 0.9, 0.2))
        )
        self.assertEqual(annotation.get("/Subtype"), "/Highlight")
        self.assertEqual(len(annotation.get("/QuadPoints")), 8)
        self.assertEqual([float(v) for v in annotation.get("/C")], [1.0, 0.9, 0.2])

    def test_build_markup_annotation_arrow_has_line_ending(self):
        annotation = build_markup_annotation(MarkupAnnotation("arrow", 10, 20, 80, 90))
        self.assertEqual(annotation.get("/Subtype"), "/Line")
        self.assertEqual([str(v) for v in annotation.get("/LE")], ["/None", "/OpenArrow"])
        self.assertEqual([float(v) for v in annotation.get("/L")], [10.0, 20.0, 80.0, 90.0])

    def test_build_markup_annotation_note_uses_contents(self):
        annotation = build_markup_annotation(
            MarkupAnnotation("note", 50, 60, 50, 60, contents="檢查這裡")
        )
        self.assertEqual(annotation.get("/Subtype"), "/Text")
        self.assertEqual(str(annotation.get("/Contents")), "檢查這裡")

    def test_build_markup_annotation_callout_has_leader_line(self):
        annotation = build_markup_annotation(
            MarkupAnnotation("callout", 40, 60, 180, 140, contents="看這裡")
        )
        self.assertEqual(annotation.get("/Subtype"), "/FreeText")
        self.assertEqual(annotation.get("/IT"), "/FreeTextCallout")
        self.assertEqual(str(annotation.get("/Contents")), "看這裡")
        self.assertEqual(len(annotation.get("/CL")), 4)

    def test_build_markup_annotation_speech_is_polygon(self):
        annotation = build_markup_annotation(
            MarkupAnnotation("speech", 30, 40, 160, 120, contents="說明")
        )
        self.assertEqual(annotation.get("/Subtype"), "/Polygon")
        self.assertGreaterEqual(len(annotation.get("/Vertices")), 14)

    def test_apply_callout_and_speech_writes_annotations(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "target.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=400)
        with source.open("wb") as stream:
            writer.write(stream)

        applied = apply_markup_annotations(
            source,
            target,
            [
                (0, MarkupAnnotation("callout", 40, 50, 180, 140, contents="A")),
                (0, MarkupAnnotation("speech", 60, 80, 200, 160, contents="B")),
            ],
        )

        self.assertEqual(applied, 2)
        annots = PdfReader(str(target)).pages[0].get("/Annots")
        self.assertGreaterEqual(len(annots), 3)

    def test_apply_markup_annotations_adds_per_page(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "target.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=400)
        writer.add_blank_page(width=300, height=400)
        with source.open("wb") as stream:
            writer.write(stream)

        markups = [
            (0, MarkupAnnotation("highlight", 10, 10, 100, 30)),
            (0, MarkupAnnotation("rect", 20, 40, 120, 90)),
            (1, MarkupAnnotation("note", 50, 60, 50, 60, contents="hi")),
        ]
        applied = apply_markup_annotations(source, target, markups)

        reader = PdfReader(str(target))
        self.assertEqual(applied, 3)
        self.assertEqual(len(reader.pages[0].get("/Annots")), 2)
        self.assertEqual(len(reader.pages[1].get("/Annots")), 1)

    def test_apply_markup_annotations_requires_items(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "target.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=400)
        with source.open("wb") as stream:
            writer.write(stream)

        with self.assertRaises(ValueError):
            apply_markup_annotations(source, target, [])

    def test_crop_pdf_pages_sets_cropbox(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "target.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=400)
        writer.add_blank_page(width=300, height=400)
        with source.open("wb") as stream:
            writer.write(stream)

        cropped = crop_pdf_pages(source, target, (50, 60, 200, 300), pages_spec="1")

        reader = PdfReader(str(target))
        self.assertEqual(cropped, 1)
        box = reader.pages[0].cropbox
        self.assertEqual(
            [float(box.left), float(box.bottom), float(box.right), float(box.top)],
            [50.0, 60.0, 200.0, 300.0],
        )
        # untouched page keeps the full media box
        self.assertEqual(float(reader.pages[1].cropbox.right), 300.0)

    def test_crop_pdf_pages_clamps_to_media_box(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "target.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=400)
        with source.open("wb") as stream:
            writer.write(stream)

        crop_pdf_pages(source, target, (-20, -20, 999, 999))

        box = PdfReader(str(target)).pages[0].cropbox
        self.assertEqual(
            [float(box.left), float(box.bottom), float(box.right), float(box.top)],
            [0.0, 0.0, 300.0, 400.0],
        )

    def test_crop_pdf_pages_rejects_invalid_rect(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "target.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=400)
        with source.open("wb") as stream:
            writer.write(stream)

        with self.assertRaises(ValueError):
            crop_pdf_pages(source, target, (200, 60, 50, 300))

    def test_parse_page_groups_splits_each_part(self):
        self.assertEqual(parse_page_groups("1-3,4-6,7", 7), [[0, 1, 2], [3, 4, 5], [6]])

    def test_split_pdf_advanced_every_n_pages(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "split.zip"
        writer = PdfWriter()
        for _ in range(5):
            writer.add_blank_page(width=200, height=200)
        with source.open("wb") as stream:
            writer.write(stream)

        count = split_pdf_advanced(source, target, "every", "2")

        self.assertEqual(count, 3)  # 2 + 2 + 1
        self.assertTrue(target.exists())

    def test_split_pdf_advanced_by_ranges(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "split.zip"
        writer = PdfWriter()
        for _ in range(6):
            writer.add_blank_page(width=200, height=200)
        with source.open("wb") as stream:
            writer.write(stream)

        count = split_pdf_advanced(source, target, "ranges", "1-2,3-6")

        self.assertEqual(count, 2)

    def test_split_pdf_advanced_rejects_bad_size(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "split.zip"
        writer = PdfWriter()
        writer.add_blank_page(width=200, height=200)
        with source.open("wb") as stream:
            writer.write(stream)

        with self.assertRaises(ValueError):
            split_pdf_advanced(source, target, "every", "0")

    def test_add_bates_numbering_stamps_each_page(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "bates.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=400)
        writer.add_blank_page(width=300, height=400)
        with source.open("wb") as stream:
            writer.write(stream)

        count = add_bates_numbering(source, target, prefix="ABC-", start=1, digits=6)

        self.assertEqual(count, 2)
        reader = PdfReader(str(target))
        first = reader.pages[0].get("/Annots")[0].get_object()
        self.assertEqual(str(first.get("/Contents")), "ABC-000001")
        second = reader.pages[1].get("/Annots")[0].get_object()
        self.assertEqual(str(second.get("/Contents")), "ABC-000002")

    def test_encrypt_pdf_with_permissions_sets_owner_password(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "protected.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=400)
        with source.open("wb") as stream:
            writer.write(stream)

        encrypt_pdf_with_permissions(
            source, target, owner_password="owner123", allow_print=False, allow_copy=False
        )

        reader = PdfReader(str(target))
        self.assertTrue(reader.is_encrypted)
        # opens without a user password
        reader.decrypt("")
        self.assertEqual(len(reader.pages), 1)

    def test_encrypt_pdf_with_permissions_requires_owner_password(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "protected.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=400)
        with source.open("wb") as stream:
            writer.write(stream)

        with self.assertRaises(ValueError):
            encrypt_pdf_with_permissions(source, target, owner_password="")

    @unittest.skipUnless(
        __import__("pdf_core").PYMUPDF_AVAILABLE,
        "PyMuPDF not installed",
    )
    def test_extract_page_text_blocks_pymupdf_includes_font_metadata(self):
        import fitz

        source = Path(self.temp_dir.name) / "pymupdf-source.pdf"
        doc = fitz.open()
        page = doc.new_page(width=300, height=200)
        page.insert_text((72, 100), "Sample Text", fontname="helv", fontsize=12)
        doc.save(str(source))
        doc.close()

        blocks = extract_page_text_blocks(source, 0)

        self.assertEqual(len(blocks), 1)
        self.assertEqual(blocks[0].text, "Sample Text")
        self.assertGreater(blocks[0].font_size, 0)
        self.assertNotEqual(blocks[0].bbox, (0.0, 0.0, 0.0, 0.0))

    def test_resolve_system_font_file_maps_noto_cjk_away_from_arial(self):
        path = resolve_system_font_file("NotoSansCJKsc-Bold", 16)
        if not path:
            self.skipTest("Windows CJK fonts not installed")
        name = Path(path).name.lower()
        self.assertNotIn("arial", name)
        self.assertTrue(any(token in name for token in ("msyh", "msjh", "mingliu", "simsun", "simhei")))

    @unittest.skipUnless(
        __import__("pdf_core").PYMUPDF_AVAILABLE,
        "PyMuPDF not installed",
    )
    def test_extract_page_text_blocks_reads_chinese(self):
        import fitz

        source = Path(self.temp_dir.name) / "cjk-source.pdf"
        doc = fitz.open()
        page = doc.new_page(width=300, height=200)
        page.insert_textbox(fitz.Rect(40, 60, 260, 120), "這個要修改", fontname="china-t", fontsize=16)
        doc.save(str(source))
        doc.close()

        blocks = extract_page_text_blocks(source, 0)
        combined = "".join(block.text for block in blocks)
        self.assertIn("這個要修改", combined.replace("\n", ""))

    @unittest.skipUnless(
        __import__("pdf_core").PYMUPDF_AVAILABLE,
        "PyMuPDF not installed",
    )
    def test_apply_erase_marks_redacts_text(self):
        import fitz

        source = Path(self.temp_dir.name) / "erase-source.pdf"
        target = Path(self.temp_dir.name) / "erase-target.pdf"
        doc = fitz.open()
        page = doc.new_page(width=300, height=200)
        page.insert_text((72, 100), "SECRET", fontname="helv", fontsize=16)
        doc.save(str(source))
        doc.close()

        apply_erase_marks(
            source,
            target,
            [
                EraseMark(
                    page_index=0,
                    kind="rect",
                    color_rgb=(1.0, 1.0, 1.0),
                    points=((60.0, 80.0), (170.0, 130.0)),
                    radius=12.0,
                )
            ],
            remove_content=True,
        )
        remaining = fitz.open(str(target))[0].get_text()
        self.assertNotIn("SECRET", remaining)

    @unittest.skipUnless(
        __import__("pdf_core").PYMUPDF_AVAILABLE,
        "PyMuPDF not installed",
    )
    def test_apply_erase_then_text_overlays_replaces_erased_text(self):
        import fitz

        source = Path(self.temp_dir.name) / "erase-text-source.pdf"
        target = Path(self.temp_dir.name) / "erase-text-target.pdf"
        doc = fitz.open()
        page = doc.new_page(width=300, height=200)
        page.insert_text((72, 100), "SECRET", fontname="helv", fontsize=16)
        doc.save(str(source))
        doc.close()

        apply_erase_then_text_overlays(
            source,
            target,
            [
                EraseMark(
                    page_index=0,
                    kind="rect",
                    color_rgb=(1.0, 1.0, 1.0),
                    points=((60.0, 80.0), (200.0, 130.0)),
                    radius=12.0,
                )
            ],
            [
                {
                    "page_index": 0,
                    "pdf_x": 72,
                    "pdf_y": 90,
                    "text": "OPEN",
                    "font_size": 12,
                    "cover": False,
                    "rect_width": 80,
                    "rect_height": 24,
                    "font_key": "helvetica",
                    "bold": False,
                    "color_rgb": (0.0, 0.0, 0.0),
                    "shape": "box",
                }
            ],
            remove_content=True,
        )
        extracted = fitz.open(str(target))[0].get_text()
        self.assertNotIn("SECRET", extracted)
        self.assertIn("OPEN", extracted)

    @unittest.skipUnless(
        __import__("pdf_core").PYMUPDF_AVAILABLE,
        "PyMuPDF not installed",
    )
    def test_replace_text_block_seamless_rewrites_text(self):
        import fitz

        source = Path(self.temp_dir.name) / "pymupdf-edit.pdf"
        target = Path(self.temp_dir.name) / "pymupdf-edited.pdf"
        doc = fitz.open()
        page = doc.new_page(width=300, height=200)
        page.insert_text((72, 100), "Old Label", fontname="helv", fontsize=12)
        page.insert_text((72, 130), "Keep", fontname="helv", fontsize=12)
        doc.save(str(source))
        doc.close()

        block = next(item for item in extract_page_text_blocks(source, 0) if item.text == "Old Label")
        replace_text_block_seamless(source, target, 0, block, "New Label")

        text = fitz.open(str(target))[0].get_text().replace("\xa0", " ")
        self.assertIn("New Label", text)
        self.assertNotIn("Old Label", text)
        self.assertIn("Keep", text)

    def _write_simple_text_pdf(self, path: Path, texts: list[str]) -> None:
        writer = PdfWriter()
        for text in texts:
            page = writer.add_blank_page(width=300, height=400)
            stream = DecodedStreamObject()
            stream.set_data(f"BT /F1 12 Tf 72 300 Td ({text}) Tj ET".encode("ascii"))
            page[NameObject("/Contents")] = stream
            page[NameObject("/Resources")] = DictionaryObject()
        with path.open("wb") as output:
            writer.write(output)

    def _write_blank_pdf(self, path: Path, page_count: int) -> None:
        writer = PdfWriter()
        for _ in range(page_count):
            writer.add_blank_page(width=300, height=400)
        with path.open("wb") as output:
            writer.write(output)

    def test_insert_pdf_pages_at_index(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        insert_from = Path(self.temp_dir.name) / "insert.pdf"
        target = Path(self.temp_dir.name) / "target.pdf"
        self._write_blank_pdf(source, 3)
        self._write_blank_pdf(insert_from, 2)

        inserted = insert_pdf_pages(source, insert_from, target, at_index=1)

        self.assertEqual(inserted, 2)
        self.assertEqual(len(PdfReader(str(target)).pages), 5)

    def test_insert_pdf_pages_with_pages_spec(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        insert_from = Path(self.temp_dir.name) / "insert.pdf"
        target = Path(self.temp_dir.name) / "target.pdf"
        self._write_blank_pdf(source, 2)
        self._write_blank_pdf(insert_from, 3)

        inserted = insert_pdf_pages(source, insert_from, target, at_index=0, pages_spec="2")

        self.assertEqual(inserted, 1)
        self.assertEqual(len(PdfReader(str(target)).pages), 3)

    def test_replace_pdf_pages_replaces_consecutive_pages(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        replacement = Path(self.temp_dir.name) / "replacement.pdf"
        target = Path(self.temp_dir.name) / "target.pdf"
        self._write_blank_pdf(source, 4)
        self._write_blank_pdf(replacement, 2)

        replaced = replace_pdf_pages(source, replacement, target, start_index=1, pages_spec="1-2")

        self.assertEqual(replaced, 2)
        self.assertEqual(len(PdfReader(str(target)).pages), 4)

    def test_replace_pdf_pages_rejects_out_of_range(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        replacement = Path(self.temp_dir.name) / "replacement.pdf"
        target = Path(self.temp_dir.name) / "target.pdf"
        self._write_blank_pdf(source, 2)
        self._write_blank_pdf(replacement, 2)

        with self.assertRaises(ValueError):
            replace_pdf_pages(source, replacement, target, start_index=1)

    def test_compare_pdf_text_reports_differences(self):
        left = Path(self.temp_dir.name) / "left.pdf"
        right = Path(self.temp_dir.name) / "right.pdf"
        report = Path(self.temp_dir.name) / "compare.txt"
        self._write_simple_text_pdf(left, ["Alpha"])
        self._write_simple_text_pdf(right, ["Beta"])

        diff_count = compare_pdf_text(left, right, report)

        self.assertEqual(diff_count, 1)
        content = report.read_text(encoding="utf-8")
        self.assertIn("--- Page 1 DIFF ---", content)
        self.assertIn("[Left]", content)
        self.assertIn("[Right]", content)

    def test_compare_pdf_text_reports_no_differences(self):
        left = Path(self.temp_dir.name) / "left.pdf"
        right = Path(self.temp_dir.name) / "right.pdf"
        report = Path(self.temp_dir.name) / "compare.txt"
        self._write_simple_text_pdf(left, ["Same"])
        self._write_simple_text_pdf(right, ["Same"])

        diff_count = compare_pdf_text(left, right, report)

        self.assertEqual(diff_count, 0)
        self.assertIn("All pages match.", report.read_text(encoding="utf-8"))

    def test_split_pdf_by_bookmarks_creates_zip_sections(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        outlined = Path(self.temp_dir.name) / "outlined.pdf"
        target = Path(self.temp_dir.name) / "bookmarks.zip"
        self._write_blank_pdf(source, 5)
        apply_outline(
            source,
            outlined,
            [
                BookmarkItem("封面", 0, 0),
                BookmarkItem("正文", 2, 0),
                BookmarkItem("附錄", 4, 0),
            ],
        )

        count = split_pdf_by_bookmarks(outlined, target)

        self.assertEqual(count, 3)
        self.assertTrue(target.exists())

    def test_split_pdf_by_bookmarks_requires_bookmarks(self):
        source = Path(self.temp_dir.name) / "source.pdf"
        target = Path(self.temp_dir.name) / "bookmarks.zip"
        self._write_blank_pdf(source, 2)

        with self.assertRaises(ValueError):
            split_pdf_by_bookmarks(source, target)

    @unittest.skipUnless(
        __import__("pdf_core").PYMUPDF_AVAILABLE,
        "PyMuPDF not installed",
    )
    def test_apply_text_markups_for_query_highlights_matches(self):
        import fitz

        source = Path(self.temp_dir.name) / "markup-source.pdf"
        target = Path(self.temp_dir.name) / "markup-target.pdf"
        doc = fitz.open()
        page = doc.new_page(width=300, height=200)
        page.insert_text((72, 100), "Hello World", fontname="helv", fontsize=12)
        doc.save(str(source))
        doc.close()

        count = apply_text_markups_for_query(source, target, "Hello", kind="highlight")

        self.assertGreaterEqual(count, 1)
        result_doc = fitz.open(str(target))
        try:
            self.assertIsNotNone(next(result_doc[0].annots(), None))
        finally:
            result_doc.close()

    @unittest.skipUnless(
        __import__("pdf_core").PYMUPDF_AVAILABLE,
        "PyMuPDF not installed",
    )
    def test_secure_redact_query_removes_text(self):
        import fitz

        source = Path(self.temp_dir.name) / "redact-source.pdf"
        target = Path(self.temp_dir.name) / "redact-target.pdf"
        doc = fitz.open()
        page = doc.new_page(width=300, height=200)
        page.insert_text((72, 100), "Secret Data", fontname="helv", fontsize=12)
        doc.save(str(source))
        doc.close()

        count = secure_redact_query(source, target, "Secret")

        self.assertGreaterEqual(count, 1)
        text = fitz.open(str(target))[0].get_text()
        self.assertNotIn("Secret", text)

    @unittest.skipUnless(
        __import__("pdf_core").PYMUPDF_AVAILABLE,
        "PyMuPDF not installed",
    )
    def test_add_text_stamp_writes_stamp_text(self):
        import fitz

        source = Path(self.temp_dir.name) / "stamp-source.pdf"
        target = Path(self.temp_dir.name) / "stamp-target.pdf"
        doc = fitz.open()
        doc.new_page(width=300, height=400)
        doc.save(str(source))
        doc.close()

        add_text_stamp(source, target, "APPROVED", 0, 72.0, 300.0, 140.0, 40.0)

        text = fitz.open(str(target))[0].get_text()
        self.assertIn("APPROVED", text)

    def test_compare_pdf_visual_reports_differences(self):
        left = Path(self.temp_dir.name) / "left.pdf"
        right = Path(self.temp_dir.name) / "right.pdf"
        report = Path(self.temp_dir.name) / "compare-visual.pdf"
        left_image = Path(self.temp_dir.name) / "left.png"
        right_image = Path(self.temp_dir.name) / "right.png"
        Image.new("RGB", (200, 240), "white").save(left_image)
        Image.new("RGB", (200, 240), "black").save(right_image)
        images_to_pdf([left_image], left)
        images_to_pdf([right_image], right)

        diff_count = compare_pdf_visual(left, right, report)

        self.assertEqual(diff_count, 1)
        self.assertTrue(report.exists())
        self.assertGreater(report.stat().st_size, 0)

    def test_compare_pdf_visual_reports_no_differences(self):
        left = Path(self.temp_dir.name) / "left.pdf"
        right = Path(self.temp_dir.name) / "right.pdf"
        report = Path(self.temp_dir.name) / "compare-visual.pdf"
        self._write_blank_pdf(left, 1)
        self._write_blank_pdf(right, 1)

        diff_count = compare_pdf_visual(left, right, report)

        self.assertEqual(diff_count, 0)
        self.assertTrue(report.exists())

    @unittest.skipUnless(
        __import__("pdf_core").PYMUPDF_AVAILABLE,
        "PyMuPDF not installed",
    )
    def test_list_and_fill_form_fields(self):
        import fitz

        source = Path(self.temp_dir.name) / "form-source.pdf"
        target = Path(self.temp_dir.name) / "form-filled.pdf"
        doc = fitz.open()
        page = doc.new_page(width=300, height=400)
        widget = fitz.Widget()
        widget.rect = fitz.Rect(50, 50, 250, 90)
        widget.field_name = "Name"
        widget.field_type = fitz.PDF_WIDGET_TYPE_TEXT
        widget.field_value = "Alice"
        widget.text_font = "Helv"
        widget.text_fontsize = 12
        page.add_widget(widget)
        doc.save(str(source))
        doc.close()

        fields = list_form_fields(source)
        self.assertEqual(len(fields), 1)
        self.assertEqual(fields[0].name, "Name")
        self.assertEqual(fields[0].field_type, "text")
        self.assertEqual(fields[0].value, "Alice")

        updated = fill_form_fields(source, target, {"Name": "Bob"})
        self.assertEqual(updated, 1)
        filled = list_form_fields(target)
        self.assertEqual(filled[0].value, "Bob")
        flattened = Path(self.temp_dir.name) / "form-flat.pdf"
        self.assertGreaterEqual(flatten_form_fields(target, flattened), 1)

    @unittest.skipUnless(
        __import__("pdf_core").PYMUPDF_AVAILABLE,
        "PyMuPDF not installed",
    )
    def test_pdf_to_docx_from_text_layer(self):
        import fitz
        from docx import Document

        source = Path(self.temp_dir.name) / "report.pdf"
        target = Path(self.temp_dir.name) / "report.docx"
        doc = fitz.open()
        page = doc.new_page(width=400, height=500)
        page.insert_text((72, 72), "Quarterly contracted sales")
        doc.save(str(source))
        doc.close()

        with patch("pdf_core._pdf_to_docx_via_word", return_value=(False, "skip")):
            with patch("pdf_core._pdf_to_office_via_libreoffice", return_value=(False, "skip")):
                with patch("pdf_core._pdf2docx_available", return_value=False):
                    pages, method = pdf_to_docx(source, target)

        self.assertEqual(pages, 1)
        self.assertEqual(method, "text")
        self.assertTrue(target.exists())
        texts = [paragraph.text for paragraph in Document(str(target)).paragraphs]
        self.assertTrue(any("Quarterly contracted sales" in text for text in texts))
        run = next(run for paragraph in Document(str(target)).paragraphs for run in paragraph.runs if run.text.strip())
        self.assertEqual(run.font.name, "Calibri")

    @unittest.skipUnless(
        __import__("pdf_core").PYMUPDF_AVAILABLE,
        "PyMuPDF not installed",
    )
    def test_pdf_to_docx_keeps_landscape_page_size(self):
        import fitz
        from docx import Document

        source = Path(self.temp_dir.name) / "slide.pdf"
        target = Path(self.temp_dir.name) / "slide.docx"
        doc = fitz.open()
        page = doc.new_page(width=960, height=540)
        page.insert_text((40, 80), "中国海外宏洋集团有限公司审计策略正文段落测试", fontname="china-t", fontsize=12)
        page.insert_text((40, 120), "根据近期会谈我们呈报审计策略概述并说明计划范围", fontname="china-t", fontsize=11)
        doc.save(str(source))
        doc.close()

        with patch("pdf_core._pdf_to_docx_via_word", return_value=(False, "skip")):
            with patch("pdf_core._pdf_to_office_via_libreoffice", return_value=(False, "skip")):
                with patch("pdf_core._convert_pdf_to_docx_pdf2docx", return_value=False):
                    pages, method = pdf_to_docx(source, target)

        self.assertEqual(method, "text")
        section = Document(str(target)).sections[0]
        self.assertGreater(int(section.page_width), int(section.page_height))
        self.assertEqual(len(Document(str(target)).tables), 0)

    @unittest.skipUnless(
        __import__("pdf_core").PYMUPDF_AVAILABLE,
        "PyMuPDF not installed",
    )
    def test_pdf_to_docx_skips_word_when_cjk_would_be_dropped(self):
        import fitz
        from docx import Document

        source = Path(self.temp_dir.name) / "cjk-report.pdf"
        target = Path(self.temp_dir.name) / "cjk-report.docx"
        doc = fitz.open()
        page = doc.new_page(width=400, height=500)
        page.insert_text(
            (36, 72),
            "金額核對工作流程審核香港上市公司年報中英文版本差異翻譯完整性",
            fontname="china-t",
            fontsize=11,
        )
        page.insert_text(
            (36, 96),
            "金額核對工作流程審核香港上市公司年報中英文版本差異翻譯完整性",
            fontname="china-t",
            fontsize=11,
        )
        doc.save(str(source))
        doc.close()

        def fake_word(_source, output):
            dropped = Document()
            dropped.add_paragraph("Skill annual-report-bilingual-audit HKFRS")
            dropped.save(str(output))
            return True, "word"

        with patch("pdf_core._pdf_to_docx_via_word", side_effect=fake_word):
            with patch("pdf_core._pdf_to_office_via_libreoffice", return_value=(False, "skip")):
                with patch("pdf_core._convert_pdf_to_docx_pdf2docx", return_value=False):
                    pages, method = pdf_to_docx(source, target)

        self.assertEqual(pages, 1)
        self.assertEqual(method, "text")
        texts = "\n".join(paragraph.text for paragraph in Document(str(target)).paragraphs)
        self.assertIn("金額核對", texts)
        self.assertGreater(count_cjk_chars(texts), 20)
        run = next(run for paragraph in Document(str(target)).paragraphs for run in paragraph.runs if "金額" in run.text)
        from docx.oxml.ns import qn

        self.assertEqual(run.font.name, "微軟正黑體")
        self.assertEqual(run._element.rPr.rFonts.get(qn("w:eastAsia")), "微軟正黑體")

    @unittest.skipUnless(
        __import__("pdf_core").PYMUPDF_AVAILABLE,
        "PyMuPDF not installed",
    )
    def test_pdf_to_docx_text_layer_keeps_bilingual_columns(self):
        import fitz
        from docx import Document

        source = Path(self.temp_dir.name) / "bilingual.pdf"
        target = Path(self.temp_dir.name) / "bilingual.docx"
        doc = fitz.open()
        page = doc.new_page(width=500, height=420)
        pairs = [
            (80, "BUSINESS REVIEW", "業務回顧"),
            (120, "Revenue grew", "收入上升"),
            (160, "in first half", "於上半年"),
            (200, "Outlook remains", "前景仍然"),
        ]
        for y, english, chinese in pairs:
            page.insert_text((40, y), english, fontsize=11)
            page.insert_text((300, y), chinese, fontname="china-t", fontsize=11)
        doc.save(str(source))
        doc.close()

        with patch("pdf_core._pdf_to_docx_via_word", return_value=(False, "skip")):
            with patch("pdf_core._pdf_to_office_via_libreoffice", return_value=(False, "skip")):
                with patch("pdf_core._convert_pdf_to_docx_pdf2docx", return_value=False):
                    pages, method = pdf_to_docx(source, target)

        self.assertEqual(pages, 1)
        self.assertEqual(method, "text")
        document = Document(str(target))
        self.assertGreaterEqual(len(document.tables), 1)
        table = document.tables[0]
        left = " ".join(row.cells[0].text for row in table.rows)
        right = " ".join(row.cells[1].text for row in table.rows)
        self.assertIn("BUSINESS", left)
        self.assertIn("業務回顧", right)
        self.assertIn("Revenue", left)
        self.assertIn("收入", right)

    @unittest.skipUnless(
        __import__("pdf_core").PYMUPDF_AVAILABLE,
        "PyMuPDF not installed",
    )
    def test_pdf_to_docx_text_layer_rebuilds_data_table(self):
        import fitz
        from docx import Document

        source = Path(self.temp_dir.name) / "table-report.pdf"
        target = Path(self.temp_dir.name) / "table-report.docx"
        doc = fitz.open()
        page = doc.new_page(width=500, height=400)
        rows = [
            (80, "Item", "Amount", "Note"),
            (120, "Rent", "1000", "A"),
            (160, "Tax", "200", "B"),
            (200, "Total", "1200", "C"),
        ]
        for y, first, second, third in rows:
            page.insert_text((50, y), first, fontsize=11)
            page.insert_text((220, y), second, fontsize=11)
            page.insert_text((360, y), third, fontsize=11)
        doc.save(str(source))
        doc.close()

        with patch("pdf_core._pdf_to_docx_via_word", return_value=(False, "skip")):
            with patch("pdf_core._pdf_to_office_via_libreoffice", return_value=(False, "skip")):
                with patch("pdf_core._convert_pdf_to_docx_pdf2docx", return_value=False):
                    pages, method = pdf_to_docx(source, target)

        self.assertEqual(method, "text")
        document = Document(str(target))
        self.assertGreaterEqual(len(document.tables), 1)
        joined = [" ".join(cell.text for cell in row.cells) for row in document.tables[0].rows]
        self.assertTrue(any("Item" in line and "Amount" in line for line in joined))
        self.assertTrue(any("Rent" in line and "1000" in line for line in joined))

    def test_docx_body_font_name_picks_simplified_and_traditional(self):
        self.assertEqual(_docx_body_font_name("我们的审计方法财务报表测试"), "微软雅黑")
        self.assertEqual(_docx_body_font_name("這份報告會與財務報表核對並請審閱"), "微軟正黑體")
        self.assertEqual(_docx_body_font_name("Quarterly report"), "Calibri")

    def test_ocr_filters_garbage_and_keeps_bilingual_columns(self):
        self.assertTrue(_is_noisy_ocr_token("$5!", 80))
        self.assertTrue(_is_noisy_ocr_token("BCDFGHJ", 90))
        self.assertFalse(_is_noisy_ocr_token("HKFRS", 90))
        self.assertFalse(_is_noisy_ocr_token("管理層", 80))
        words = [
            (20, 40, 80, 54, "BUSINESS", 0, 0, 0),
            (90, 40, 150, 54, "REVIEW", 0, 0, 1),
            (320, 40, 380, 54, "業務回顧", 0, 0, 2),
            (20, 80, 70, 94, "Revenue", 0, 1, 0),
            (320, 80, 370, 94, "收入", 0, 1, 1),
        ]
        text = _ocr_words_to_plain_text(words, 500)
        self.assertIn("BUSINESS", text)
        self.assertIn("業務回顧", text)
        self.assertIn("|", text)

    def test_ocr_words_to_rows_pairs_left_english_right_chinese(self):
        words = [
            (20, 40, 80, 54, "BUSINESS", 0, 0, 0),
            (90, 40, 150, 54, "REVIEW", 0, 0, 1),
            (320, 40, 400, 54, "業務回顧", 0, 0, 2),
            (20, 80, 90, 94, "Revenue", 0, 1, 0),
            (320, 80, 370, 94, "收入", 0, 1, 1),
            (20, 120, 70, 134, "grew", 0, 2, 0),
            (20, 160, 100, 174, "steadily", 0, 3, 0),
            (320, 120, 400, 134, "穩步增長", 0, 2, 1),
            (20, 200, 80, 214, "Outlook", 0, 4, 0),
            (320, 200, 380, 214, "前景", 0, 4, 1),
            (20, 240, 90, 254, "remains", 0, 5, 0),
            (320, 240, 400, 254, "仍然穩健", 0, 5, 1),
        ]
        rows = _ocr_words_to_rows(words, 500, split_at=250)
        self.assertGreaterEqual(len(rows), 3)
        self.assertEqual(rows[0][0], "BUSINESS REVIEW")
        self.assertEqual(rows[0][1], "業務回顧")
        self.assertTrue(any(row[0] == "Revenue" and row[1] == "收入" for row in rows))

    def test_detect_page_gutter_x_finds_center_valley(self):
        image = Image.new("L", (400, 300), 255)
        pixels = image.load()
        for x in range(30, 160):
            for y in range(40, 260):
                pixels[x, y] = 20
        for x in range(240, 370):
            for y in range(40, 260):
                pixels[x, y] = 20
        gutter = _detect_page_gutter_x(image)
        self.assertIsNotNone(gutter)
        self.assertGreater(gutter, 160)
        self.assertLess(gutter, 240)

    def test_detect_page_gutter_x_ignores_gray_paper_noise(self):
        image = Image.new("L", (400, 300), 210)
        pixels = image.load()
        for x in range(30, 160):
            for y in range(40, 260):
                pixels[x, y] = 40
        for x in range(240, 370):
            for y in range(40, 260):
                pixels[x, y] = 40
        gutter = _detect_page_gutter_x(image)
        self.assertIsNotNone(gutter)
        self.assertGreater(gutter, 160)
        self.assertLess(gutter, 240)

    def test_ocr_page_keeps_two_columns_for_short_headings(self):
        from docx import Document

        document = Document()
        rows = [
            ["MANAGEMENT DISCUSSION AND", "管理層討論及分析"],
            ["BUSINESS REVIEW", "業務回顧"],
        ]
        _add_docx_ocr_page(document, rows, "Calibri")
        self.assertEqual(len(document.tables), 1)
        self.assertEqual(document.tables[0].cell(0, 0).text.strip(), "MANAGEMENT DISCUSSION AND")
        self.assertEqual(document.tables[0].cell(0, 1).text.strip(), "管理層討論及分析")
        mixed = [paragraph.text for paragraph in document.paragraphs if "MANAGEMENT" in paragraph.text and "管理層" in paragraph.text]
        self.assertEqual(mixed, [])

    def test_docx_ocr_page_writes_bilingual_table(self):
        from docx import Document
        from docx.oxml.ns import qn

        document = Document()
        rows = [
            ["BUSINESS REVIEW", "業務回顧"],
            ["Revenue increased", "收入上升"],
            ["in the first half", "於上半年"],
            ["Outlook remains", "前景仍然"],
        ]
        _add_docx_ocr_page(document, rows, "Calibri")
        self.assertEqual(len(document.tables), 1)
        table = document.tables[0]
        self.assertEqual(table.cell(0, 0).text.strip(), "BUSINESS REVIEW")
        self.assertEqual(table.cell(0, 1).text.strip(), "業務回顧")
        self.assertEqual(table.cell(1, 1).text.strip(), "收入上升")
        right_run = table.cell(0, 1).paragraphs[0].runs[0]
        self.assertEqual(right_run._element.rPr.rFonts.get(qn("w:eastAsia")), "微軟正黑體")

    def test_toc_page_numbers_are_not_side_by_side_layout(self):
        rows = [
            ["审计目标及管理层的责任", "04"],
            ["审计团队", "07"],
            ["我们的审计方法", "10"],
            ["对舞弊风险的认识", "24"],
        ]
        self.assertTrue(_page_rows_are_two_column(rows))
        self.assertFalse(_rows_are_side_by_side_layout(rows))

    def test_docx_layout_table_fits_landscape_page(self):
        from docx import Document

        document = Document()
        _configure_docx_section_from_pdf(document, 960, 540)
        section = document.sections[0]
        self.assertGreater(int(section.page_width), int(section.page_height))
        rows = [
            ["BUSINESS REVIEW", "業務回顧"],
            ["Revenue increased", "收入上升"],
            ["in the first half", "於上半年"],
            ["Outlook remains", "前景仍然"],
        ]
        _add_docx_layout_table(document, rows, "Calibri", 11.0, bordered=False)
        table = document.tables[0]
        usable = int(section.page_width - section.left_margin - section.right_margin)
        cell_w = int(table.cell(0, 0).width)
        self.assertGreater(cell_w, usable * 0.3)
        self.assertLess(cell_w, usable * 0.7)
        self.assertLess(cell_w * 2, usable + 200000)

    def test_column_ocr_uses_english_first_on_left(self):
        self.assertEqual(_column_tess_lang("chi_tra+eng", prefer_cjk=False), "eng+chi_tra")
        self.assertEqual(_column_tess_lang("chi_tra+eng", prefer_cjk=True), "chi_tra+eng")
        self.assertIn("--psm 6", _tesseract_column_config("--oem 1 --psm 3 -c preserve_interword_spaces=1"))

    def test_words_to_table_rows_splits_aligned_columns(self):
        words = [
            (50, 80, 90, 92, "Item", 0, 0, 0),
            (280, 80, 340, 92, "Amount", 0, 0, 1),
            (50, 120, 90, 132, "Rent", 0, 1, 0),
            (280, 120, 320, 132, "1000", 0, 1, 1),
            (50, 160, 80, 172, "Tax", 0, 2, 0),
            (280, 160, 310, 172, "200", 0, 2, 1),
            (50, 200, 90, 212, "Total", 0, 3, 0),
            (280, 200, 320, 212, "1200", 0, 3, 1),
        ]
        rows = _words_to_table_rows(words, 500)
        self.assertEqual(len(rows), 4)
        self.assertEqual(rows[0][:2], ["Item", "Amount"])
        self.assertEqual(rows[1][:2], ["Rent", "1000"])
        self.assertTrue(_is_useful_table(rows))
        self.assertFalse(_is_useful_table([["logo"], [""]]))
        self.assertFalse(_is_useful_table([["only one column"], ["still one"]]))

    def test_table_looks_fragmented_rejects_split_words(self):
        fragmented = [
            ["HKFRS Ac", "counting S", "tandards", "Update No", ".", "3"],
            ["ection I. Am", "ended Sta", "ndards issued", "that are ap", "plicable to", "2027"],
            ["HKFRS 9 a", "nd HKFRS", "7 Classifi", "cation and", "Measure", "ment"],
        ]
        self.assertTrue(_table_looks_fragmented(fragmented))
        intact = [
            ["Standards affected", "Amendments relate to", "Members' Handbook"],
            ["HKFRS 9 and HKFRS 7", "Classification and Measurement of Financial Instruments", "Update No. 315"],
            ["HKFRS 1, HKFRS 7", "Annual Improvements to HKFRS Accounting Standards", "Update No. 316"],
        ]
        self.assertFalse(_table_looks_fragmented(intact))

    @unittest.skipUnless(
        __import__("pdf_core").PYMUPDF_AVAILABLE,
        "PyMuPDF not installed",
    )
    def test_pymupdf_page_tables_uses_default_lines_only(self):
        import fitz
        from types import SimpleNamespace

        calls: list[dict] = []

        def fake_find_tables(*_args, **kwargs):
            calls.append(kwargs)
            return SimpleNamespace(tables=[])

        source = Path(self.temp_dir.name) / "lines-only.pdf"
        doc = fitz.open()
        page = doc.new_page(width=400, height=300)
        page.insert_text((50, 80), "Item")
        page.insert_text((220, 80), "Amount")
        doc.save(str(source))
        doc.close()

        document = fitz.open(str(source))
        try:
            with patch.object(type(document[0]), "find_tables", fake_find_tables):
                rows = _pymupdf_page_tables(document[0])
        finally:
            document.close()

        self.assertEqual(rows, [])
        self.assertEqual(len(calls), 1)
        self.assertNotIn("strategy", calls[0])
        self.assertNotEqual(calls[0].get("vertical_strategy"), "text")

    @unittest.skipUnless(
        __import__("pdf_core").PYMUPDF_AVAILABLE,
        "PyMuPDF not installed",
    )
    def test_pdf_to_xlsx_reports_page_progress(self):
        import fitz

        source = Path(self.temp_dir.name) / "progress.pdf"
        target = Path(self.temp_dir.name) / "progress.xlsx"
        doc = fitz.open()
        for index in range(2):
            page = doc.new_page(width=400, height=300)
            page.insert_text((50, 80), f"Row {index + 1}")
            page.insert_text((220, 80), str((index + 1) * 10))
        doc.save(str(source))
        doc.close()

        seen: list[tuple[int, int, str]] = []

        def progress(current: int, total: int, text: str) -> None:
            seen.append((current, total, text))

        pages, method = pdf_to_xlsx(source, target, progress=progress)

        self.assertEqual(pages, 2)
        self.assertTrue(target.exists())
        self.assertTrue(any("第 1 / 2 頁" in text for _current, _total, text in seen))
        self.assertTrue(any(current == 1 and total == 2 for current, total, _text in seen))

    @unittest.skipUnless(
        __import__("pdf_core").PYMUPDF_AVAILABLE,
        "PyMuPDF not installed",
    )
    def test_pdf_to_xlsx_writes_text_sheet(self):
        import fitz
        from openpyxl import load_workbook

        source = Path(self.temp_dir.name) / "figures.pdf"
        target = Path(self.temp_dir.name) / "figures.xlsx"
        doc = fitz.open()
        page = doc.new_page(width=400, height=500)
        page.insert_text((72, 72), "Alpha 100")
        page.insert_text((72, 96), "Beta 200")
        doc.save(str(source))
        doc.close()

        pages, method = pdf_to_xlsx(source, target)

        self.assertEqual(pages, 1)
        self.assertIn(method, {"tables", "columns", "text"})
        self.assertTrue(target.exists())
        values = [cell.value for row in load_workbook(target).active.iter_rows() for cell in row]
        self.assertTrue(any(value and "Alpha" in str(value) for value in values))

    @unittest.skipUnless(
        __import__("pdf_core").PYMUPDF_AVAILABLE,
        "PyMuPDF not installed",
    )
    def test_pdf_to_xlsx_extracts_borderless_columns(self):
        import fitz
        from openpyxl import load_workbook

        source = Path(self.temp_dir.name) / "borderless.pdf"
        target = Path(self.temp_dir.name) / "borderless.xlsx"
        doc = fitz.open()
        page = doc.new_page(width=500, height=400)
        page.insert_text((50, 80), "Item")
        page.insert_text((280, 80), "Amount")
        page.insert_text((50, 120), "Rent")
        page.insert_text((280, 120), "1000")
        page.insert_text((50, 160), "Tax")
        page.insert_text((280, 160), "200")
        page.insert_text((50, 200), "Total")
        page.insert_text((280, 200), "1200")
        doc.save(str(source))
        doc.close()

        pages, method = pdf_to_xlsx(source, target)

        self.assertEqual(pages, 1)
        self.assertIn(method, {"tables", "columns"})
        rows = [tuple(cell.value for cell in row) for row in load_workbook(target).active.iter_rows()]
        joined = [" ".join("" if cell is None else str(cell) for cell in row) for row in rows]
        self.assertTrue(any("Item" in line and "Amount" in line for line in joined))
        self.assertTrue(any("Rent" in line and "1000" in line for line in joined))
        widths = {len(row) for row in rows}
        self.assertTrue(any(width >= 2 for width in widths))

    @unittest.skipUnless(
        __import__("pdf_core").PYMUPDF_AVAILABLE,
        "PyMuPDF not installed",
    )
    def test_pdf_to_xlsx_skips_junk_tables_and_keeps_columns(self):
        import fitz
        from openpyxl import load_workbook
        from types import SimpleNamespace

        source = Path(self.temp_dir.name) / "junk-table.pdf"
        target = Path(self.temp_dir.name) / "junk-table.xlsx"
        doc = fitz.open()
        page = doc.new_page(width=500, height=400)
        page.insert_text((50, 80), "Account")
        page.insert_text((280, 80), "Balance")
        page.insert_text((50, 120), "Cash")
        page.insert_text((280, 120), "500")
        page.insert_text((50, 160), "Debt")
        page.insert_text((280, 160), "80")
        doc.save(str(source))
        doc.close()

        junk = SimpleNamespace(extract=lambda: [["logo"], [""]])
        finder = SimpleNamespace(tables=[junk])
        with patch("fitz.Page.find_tables", return_value=finder):
            pages, method = pdf_to_xlsx(source, target)

        self.assertEqual(pages, 1)
        self.assertEqual(method, "columns")
        rows = [tuple(cell.value for cell in row) for row in load_workbook(target).active.iter_rows()]
        joined = [" ".join("" if cell is None else str(cell) for cell in row) for row in rows]
        self.assertTrue(any("Cash" in line and "500" in line for line in joined))

    @unittest.skipUnless(
        __import__("pdf_core").PYMUPDF_AVAILABLE,
        "PyMuPDF not installed",
    )
    def test_pdf_to_xlsx_ocr_rebuilds_columns_from_boxes(self):
        import fitz
        from openpyxl import load_workbook
        from types import SimpleNamespace

        source = Path(self.temp_dir.name) / "scan-table.pdf"
        target = Path(self.temp_dir.name) / "scan-table.xlsx"
        image_path = Path(self.temp_dir.name) / "scan.png"
        Image.new("RGB", (400, 300), "white").save(image_path)
        doc = fitz.open()
        page = doc.new_page(width=400, height=300)
        page.insert_image(page.rect, filename=str(image_path))
        doc.save(str(source))
        doc.close()

        data = {
            "text": ["Item", "Amount", "Rent", "1000", "Tax", "200"],
            "conf": ["90"] * 6,
            "left": [50, 280, 50, 280, 50, 280],
            "top": [80, 80, 120, 120, 160, 160],
            "width": [40, 60, 40, 40, 30, 30],
            "height": [12] * 6,
        }
        ocr = Mock()
        ocr.Output = SimpleNamespace(DICT="dict")
        ocr.image_to_data = Mock(return_value=data)
        ocr.image_to_string = Mock(return_value="Item Amount")
        with patch("pdf_core.ensure_ocr_available"), patch(
            "pdf_core.iter_pdf_page_images",
            return_value=[(0, Image.new("RGB", (400, 300), "white"))],
        ), patch("pdf_core.pytesseract", ocr):
            pages, method = pdf_to_xlsx(source, target)

        self.assertEqual(pages, 1)
        self.assertEqual(method, "ocr")
        rows = [tuple(cell.value for cell in row) for row in load_workbook(target).active.iter_rows()]
        joined = [" ".join("" if cell is None else str(cell) for cell in row) for row in rows]
        self.assertTrue(any("Item" in line and "Amount" in line for line in joined))
        self.assertTrue(any("Rent" in line and "1000" in line for line in joined))

    def setUp(self):
        import tempfile

        self.temp_dir = tempfile.TemporaryDirectory()

    def tearDown(self):
        self.temp_dir.cleanup()


if __name__ == "__main__":
    unittest.main()
